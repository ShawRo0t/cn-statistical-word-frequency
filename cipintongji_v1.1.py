import tkinter
import tkinter.filedialog
import tkinter.messagebox
from collections import Counter
import jieba
import docx
import os

top = tkinter.Tk()
top.geometry("500x150+400+200")
top.title('中文词频统计工具 Power By ShawRoot')
frame = tkinter.Frame(top)
# 设置外边距
frame.pack(padx=10,pady=10)
frame1 = tkinter.Frame(top)
frame1.pack(padx=10,pady=10)
frame2 = tkinter.Frame(top)
frame2.pack(padx=10,pady=10)
v = tkinter.StringVar()
v2 = tkinter.StringVar()
#设置停用和过滤词
stop_word = [' ','/',':','。','.','?','？','：','!','-','@','$','*','(',')','+','%','^','&','！','“','\'','"','”','…','[',']',',','，','=','~','《','》','>','<','；',';','、','（','）','的','	','\\','{','}','|','#','	','_',' 	']
#文件夹词频总数
all_data = {}
#载入自定义词典
try:
    jieba.load_userdict(r"Dict.txt")
except:
    tkinter.messagebox.showinfo('提示','未找到自定义词典Dict.txt！')
# 打开文件函数
def normaldata(filename,tail):
    try:
        file = open(filename,encoding='utf-8').read()
    except:
        file = open(filename,encoding='gbk',errors="ignore").read()
    words = list(jieba.cut(file))
    data = dict(Counter(words).most_common(50))
    
    resultpath = "fenci_output"
    if not os.path.exists(resultpath):
        os.mkdir(resultpath)
    result = open("fenci_output/"+tail+"分词统计结果.txt",'w',encoding='utf-8')
    for k,p in data.items():
        if k in stop_word:
            continue
        else:
            result.write("%s\t%d\n" % (k,p))
    tkinter.messagebox.showinfo('提示','分词成功！结果已生成在“fenci_output”目录下！')
def pathnormaldata(pathname,filepathname):
    try:
        file = open(pathname+'/'+filepathname,encoding='utf-8').read()
    except:
        file = open(pathname+'/'+filepathname,encoding='gbk',errors="ignore").read()
    words = list(jieba.cut(file))
    data = dict(Counter(words).most_common(50))
    for e3,y3 in data.items():
        if e3 in all_data:
            all_data[e3] += y3
        else:
            if e3 in stop_word:
                pass
            else:
                all_data[e3] = y3
    resultpath = "fenci_output"
    if not os.path.exists(resultpath):
        os.mkdir(resultpath)
    result = open("fenci_output/"+filepathname+"分词统计结果.txt",'w',encoding='utf-8')
    for k,p in data.items():
        if k in stop_word:
            continue
        else:
            result.write("%s\t%d\n" % (k,p))
def run():
    filename = v.get()
    pathname = v2.get()
    if filename and pathname:
        tkinter.messagebox.showinfo('提示','抱歉！请指定文件/文件夹中的一项！')
    elif filename:
        head, tail = os.path.split(filename)
        word_word = ""
        if(".txt" in filename):
            head, tail = os.path.split(filename)
            normaldata(filename,tail)
        elif(".doc" in filename):
            file = docx.Document(filename)
            for para in file.paragraphs:
                word_word = word_word + para.text
            word_words = list(jieba.cut(word_word))
            data = dict(Counter(word_words).most_common(50))
            resultpath = "fenci_output"
            if not os.path.exists(resultpath):
                os.mkdir(resultpath)
            result = open("fenci_output/"+tail+"分词统计结果.txt",'w',encoding='utf-8')
            for k,p in data.items():
                if k in stop_word:
                    continue
                else:
                    result.write("%s\t%d\n" % (k,p))
            tkinter.messagebox.showinfo('提示','分词成功！结果已生成在“fenci_output”目录下！')
        elif(".html" in filename):
            head, tail = os.path.split(filename)
            normaldata(filename,tail)
        else:
            tkinter.messagebox.showinfo('提示','抱歉！不是指定类型的文件')
    elif pathname:
        filepathnames = os.listdir(pathname)
        pathword_word = ""
        for filepathname in filepathnames:
            print(str(filepathname))
            if(".txt" in filepathname):
                print("[*] 正在分词："+ filepathname)
                pathnormaldata(pathname,filepathname)
            elif(".doc" in filepathname):
                print("[*] 正在分词："+ filepathname)
                try:
                    file = docx.Document(pathname+'/'+filepathname)
                except:
                    tkinter.messagebox.showinfo('提示',filepathname+'分词失败！错误：未知原因')
                    continue
                for para in file.paragraphs:
                    pathword_word = pathword_word + para.text
                pathword_words = list(jieba.cut(pathword_word))
                data = dict(Counter(pathword_words).most_common(50))
                for e1,y1 in data.items():
                    if e1 in all_data:
                        all_data[e1] += y1
                    else:
                        if e1 in stop_word:
                            pass
                        else:
                            all_data[e1] = y1
                resultpath = "fenci_output"
                if not os.path.exists(resultpath):
                    os.mkdir(resultpath)
                result = open("fenci_output/"+filepathname+"分词统计结果.txt",'w',encoding='utf-8')
                for e2,y2 in data.items():
                    if e2 in stop_word:
                        continue
                    else:
                        result.write("%s\t%d\n" % (e2,y2))
            elif(".html" in filepathname):
                print("[*] 正在分词："+ filepathname)
                pathnormaldata(pathname,filepathname)
            else:
                pass
        all_result = open("fenci_output/"+"文件夹总分词统计结果.txt",'w',encoding='utf-8')
        for e4,y4 in all_data.items():
            all_result.write("%s\t%d\n" % (e4,y4))
        tkinter.messagebox.showinfo('提示','文件夹分词成功！结果已生成在“fenci_output”目录下！')
    else:
        tkinter.messagebox.showinfo('提示','请载入要分词的文件或文件夹！')
def fileopen():
    # 清除用户的输入项
    v.set('')
    filenamepath = tkinter.filedialog.askopenfilename()
    if(filenamepath):
        v.set(filenamepath)
def fileopenpath():
    # 清除用户的输入项
    v2.set('')
    filepath = tkinter.filedialog.askdirectory()
    if(filepath):
        v2.set(filepath)
file_path = tkinter.Entry(frame,width=50,textvariable=v).pack(fill=tkinter.X,side=tkinter.LEFT)
button1 = tkinter.Button(frame,width=20,text='选择文件',activeforeground='white',activebackground='RoyalBlue',command=fileopen).pack(fill=tkinter.X,padx=10)
file_path = tkinter.Entry(frame1,width=50,textvariable=v2).pack(fill=tkinter.X,side=tkinter.LEFT)
button1 = tkinter.Button(frame1,width=20,text='选择文件夹',activeforeground='white',activebackground='RoyalBlue',command=fileopenpath).pack(fill=tkinter.X,padx=10)
button2 = tkinter.Button(frame2,width=20,text='统计',activeforeground='white',activebackground='RoyalBlue',command=run).pack(fill=tkinter.X,side=tkinter.LEFT)
button3 = tkinter.Button(frame2,width=20,text='退出',activeforeground='white',activebackground='RoyalBlue',command=top.quit).pack(fill=tkinter.X,padx=10)
top.mainloop()
